
from docx.shared import Inches
from docx import Document


def report_content(dict):
    
    #get infor details from dictionary
    name = dict["name"]
    company = dict["company"]
    person_debt = dict["person_debt"]
    business_debt= dict["business_debt"]
    person_income= dict["person_income"]
    person_year1= dict["person_year1"]
    person_year2= dict["person_year2"]
    person_year3= dict["person_year3"]
    business_income= dict["business_income"]
    business_year1= dict["business_year1"]
    business_year2= dict["business_year2"]
    business_year3= dict["business_year3"]
    DTO= dict["DTO"]
    DTI= dict["DTI"]
    GDSCR= dict["GDSCR"]
    capacity= dict["capacity"]
    balance= dict["balance"]
    ex= dict["ex"]
    eq = dict["eq"]
    tu = dict["tu"]

    #Calculation
    try:
        total_income = int(person_income) + int(business_income)
    except ValueError:
        total_income = 0
    try:

        total_debt_pmt = int(person_debt) + int(business_debt)
    except ValueError:
        total_debt_pmt = 0

    #Change to string to be able to concate in the documents
    total_income = str(total_income)
    total_debt_pmt = str(total_debt_pmt)
    GDSCR = str(GDSCR)
    DTO = str(DTO)
    DTI = str(DTI)
    balance = str(balance)
    capacity= str(capacity)

    #Initiate document
    document = Document()

    #Header
    document.add_picture('picture/wcf_logo.png', width=Inches(3))
    document.add_heading('Personal Business Financial Analysis', 0)
    #document.add_paragraph('for '+ name)
    document.add_paragraph('Prepared by Wallace Capital Funding, LLC ')

    #Person Name

    #Business Name

    #Personal Credit Profile
    document.add_heading('Personal Credit Profile', level=1)
    document.add_paragraph('Credit Scores', style='List Bullet')
    records = (
        (name, ex, eq, tu)
    )
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'TU'
    hdr_cells[2].text = 'EQ'
    hdr_cells[3].text = 'EX'

    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = ex
    row_cells[2].text = eq
    row_cells[3].text = tu

    """
      for Name, TU, EQ, EX in records:
        row_cells = table.add_row().cells
        row_cells[0].text = Name
        row_cells[1].text = TU
        row_cells[2].text = EQ
        row_cells[3].text = EX
    
    """
  

    document.add_paragraph('PCW: '+ name, style='List Bullet')
    document.add_paragraph('   Debt-to-Owing Ratio (DTO): '+DTO)
    document.add_paragraph('   Debt-to-Income Ratio (DTI): '+DTI)
    
    #Cashflow analysis
    document.add_heading('Cash Flow Analysis ', level=1)
    document.add_paragraph('GDSCR', style='List Bullet')
    document.add_paragraph("   The Global Cash Flow Sheet allows us to determine the total operating income and total debt. We compute the operating income using information from the company's balance sheet, profit and loss statement, and tax returns. The total debt is calculated based on recent credit report statements. The average net operating income is $" + total_income + " per month. The total monthly debt payments is $"+ total_debt_pmt +". This results in a Gross Debt Service Coverage Ratio (GDSCR) of " + GDSCR + ".")
    
    #Bank Statement Analsyis
    document.add_heading('Bank Statement Analysis', level=1)
    document.add_paragraph("   We conducted a bank statement analysis and the analysis showed that the average monthly balance is "+ balance +". While based on your average deposit and withdrawal, you have a capacity of "+ capacity +"." , style='List Bullet')
    
    #Recommendations
    document.add_heading('Recommendation', level=1)
    document.add_paragraph("   Write recomendation here", style='List Bullet')
    
    return document
#document.add_page_break()
#document.save('\documents\demo.docx')

import os
import tempfile

def tmp_dir():
    with tempfile.TemporaryDirectory() as tmpdirname:
        return(tempfile.gettempdir())

def generate_report(dict):

    #get name
    name = dict["name"]

    #Created temporary directory
    directory = tmp_dir()

    filename = "BFA for " + name + ".docx"
    file_path = os.path.join(directory, filename)

    #Generate Report
    report = report_content(dict)

    #Save the report in temporary directory
    report.save(file_path)
    
    #file = open(file_path, "w")
    #file.write("you have realy good GDSCR "+ docName)
    #file.close()
    return file_path

